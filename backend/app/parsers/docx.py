"""DOCX parser: paragraphs + table cells (python-docx) -> one event per non-empty line."""
from __future__ import annotations

import io
from typing import Iterable, Iterator

from .base import BaseParser, ParsedEvent
from .tabular import line_event


def available() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except Exception:
        return False


class DocxParser(BaseParser):
    name = "Word document (DOCX)"
    family = "document.docx"
    binary = True
    extensions = (".docx", ".docm")

    def sniff(self, sample: list[str], filename: str = "", head: bytes = b"") -> float:
        if filename.lower().endswith(self.extensions):
            return 1.0 if head.startswith(b"PK\x03\x04") else 0.7
        return 0.0

    def parse(self, lines: Iterable[str]) -> Iterator[ParsedEvent]:
        for i, line in enumerate(lines):
            if line.strip():
                yield line_event(line.strip(), {"line_no": str(i + 1)})

    def parse_bytes(self, data: bytes) -> Iterator[ParsedEvent]:
        try:
            import docx
        except Exception as exc:
            raise RuntimeError(f"python-docx not installed: {exc}")
        doc = docx.Document(io.BytesIO(data))
        n = 0
        for para in doc.paragraphs:
            for line in (para.text or "").splitlines():
                s = line.strip()
                if not s:
                    continue
                n += 1
                extra = {"line_no": str(n), "kind": "paragraph"}
                try:
                    style = para.style.name if para.style is not None else ""
                except Exception:
                    style = ""
                if style and style != "Normal":
                    extra["style"] = style
                yield line_event(s, extra)
        for ti, table in enumerate(doc.tables, start=1):
            for ri, row in enumerate(table.rows, start=1):
                cells = [c.text.strip() for c in row.cells]
                joined = " | ".join(c for c in cells if c)
                if not joined:
                    continue
                n += 1
                yield line_event(joined, {"line_no": str(n), "kind": "table", "table": str(ti), "row": str(ri)})
